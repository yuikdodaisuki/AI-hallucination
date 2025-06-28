"""数据读取和整合功能"""
import os
import json
from src.AIquest.config import (
    DATA_SOURCES, OUTPUT_CONFIG, METRIC_DATA_MAPPING, REQUIRED_DIRECTORIES,
    is_school_extraction_enabled, get_school_extraction_config, 
    get_traditional_extraction_config, get_attachment_config,
    get_consolidated_dir_path,get_consolidated_dir_name,get_output_config
)


class DataReader:
    """数据读取和整合功能"""
    
    def __init__(self):
        # 🔥 修正路径计算 - data目录和AIquest同级 🔥
        self.current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # src/AIquest
        self.src_dir = os.path.dirname(self.current_dir)  # src
        self.data_dir = os.path.join(self.src_dir, 'data')  # src/data
        self._ensure_directories_exist()
        
        print(f"📁 DataReader路径信息:")
        print(f"  AIquest目录: {self.current_dir}")
        print(f"  src目录: {self.src_dir}")
        print(f"  data目录: {self.data_dir}")

        # 🔥 新增：显示当前使用的整合目录 🔥
        current_consolidated_dir = get_consolidated_dir_path(self.data_dir)
        mode = "智能截取模式" if is_school_extraction_enabled() else "传统模式"
        print(f"  📂 当前整合目录: {current_consolidated_dir} ({mode})")
    
    def _ensure_directories_exist(self):
        """确保所有必需的目录存在（跳过已存在的moepolicies）"""
        for directory in REQUIRED_DIRECTORIES:
            dir_path = os.path.join(self.data_dir, directory)
            if not os.path.exists(dir_path):
                try:
                    os.makedirs(dir_path, exist_ok=True)
                    print(f"  📁 创建目录: {dir_path}")
                    
                    # 为新创建的目录添加README文件
                    readme_path = os.path.join(dir_path, 'README.md')
                    readme_content = self._generate_readme_content(directory)
                    with open(readme_path, 'w', encoding='utf-8') as f:
                        f.write(readme_content)
                        
                except Exception as e:
                    print(f"  ❌ 创建目录失败 {dir_path}: {e}")
    
    def _generate_readme_content(self, dir_name):
        """生成目录说明文档"""
        descriptions = {
            'esi_subjects': 'ESI学科相关数据，包含进入ESI排名的学科信息',
            'esi_subjects/esi_top1percent': 'ESI前1%学科专门数据',
            'esi_subjects/esi_top1permille': 'ESI前1‰学科专门数据',
            'ruanke_subjects': '软科中国最好学科排名相关数据',
            'subject_evaluation': '教育部学科评估相关数据，包含A+、A、A-等评级信息',
            'undergraduate_majors': '本科专业相关数据',
            'undergraduate_majors/total_majors': '本科专业总数统计数据',
            'undergraduate_majors/certified_majors': '通过专业认证的本科专业数据',
            'undergraduate_majors/national_first_class': '国家级一流本科专业建设点数据',
            'undergraduate_majors/provincial_first_class': '省级一流本科专业建设点数据',
            'consolidated': '整合后的数据文件'
        }
        
        description = descriptions.get(dir_name, '数据存储目录')
        
        return f"""# {dir_name} 数据目录

## 目录说明
{description}

## 数据格式要求
- 文件格式：JSON
- 编码：UTF-8
- 文件命名：建议使用有意义的名称，如 `university_name.json`

## 使用说明
1. 将相关数据文件放置在此目录中
2. 系统会自动扫描并读取所有 `.json` 文件
3. 数据会被自动整合到问答系统中

## 更新时间
{self._get_current_time()}
"""
    
    def _get_current_time(self):
        """获取当前时间字符串"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    def read_data_from_source(self, source_name):
        """从指定数据源读取数据"""
        if source_name not in DATA_SOURCES:
            print(f"警告: 未知的数据源: {source_name}")
            return []
        
        # 🔥 修正路径计算 - 基于src/data目录 🔥
        source_relative_path = DATA_SOURCES[source_name]
        
        # 移除路径前缀并计算正确路径
        if source_relative_path.startswith('../data/'):
            clean_path = source_relative_path.replace('../data/', '')
        else:
            clean_path = source_relative_path
        
        source_path = os.path.join(self.data_dir, clean_path)
        print(f"    正在读取数据源: {source_path}")
        
        all_data = []
        
        if not os.path.exists(source_path):
            print(f"    警告: 数据源路径不存在: {source_path}")
            # 🔥 添加调试信息 🔥
            print(f"    调试信息:")
            print(f"      配置的源路径: {source_relative_path}")
            print(f"      清理后的路径: {clean_path}")
            print(f"      data目录: {self.data_dir}")
            print(f"      最终路径: {source_path}")
            print(f"      目录是否存在: {os.path.exists(source_path)}")
            
            # 🔥 列出data目录内容 🔥
            if os.path.exists(self.data_dir):
                print(f"    data目录内容:")
                for item in os.listdir(self.data_dir):
                    item_path = os.path.join(self.data_dir, item)
                    if os.path.isdir(item_path):
                        print(f"      📁 {item}/")
                        # 如果是subject_evaluation目录，列出其内容
                        if item == 'subject_evaluation':
                            try:
                                sub_items = os.listdir(item_path)
                                json_files = [f for f in sub_items if f.endswith('.json')]
                                print(f"        📄 JSON文件: {len(json_files)} 个")
                                for json_file in json_files[:3]:
                                    print(f"          - {json_file}")
                            except Exception as e:
                                print(f"        ❌ 无法读取目录内容: {e}")
                    else:
                        print(f"      📄 {item}")
            else:
                print(f"    ❌ data目录不存在: {self.data_dir}")
            
            return all_data
        
        # 递归读取所有JSON文件
        file_count = 0
        for root, _, files in os.walk(source_path):
            for file_name in files:
                if file_name.endswith('.json') and not file_name.startswith('combined_'):
                    file_path = os.path.join(root, file_name)
                    print(f"      读取文件: {file_path}")
                    data = self._read_json_file(file_path)
                    if data:
                        normalized_data = self._normalize_json_data(data, file_path)
                        all_data.extend(normalized_data)
                        file_count += 1
        
        print(f"    从 {source_path} 读取到 {len(all_data)} 条数据，处理了 {file_count} 个JSON文件")
        return all_data
    
    def _read_json_file(self, file_path):
        """读取单个JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except json.JSONDecodeError as e:
            print(f"      警告: {file_path} 不是有效的JSON文件: {e}")
            return None
        except Exception as e:
            print(f"      错误: 读取 {file_path} 失败: {e}")
            return None
    
    def _normalize_json_data(self, data, file_path):
        """标准化JSON数据格式，包含附件内容提取"""
        if isinstance(data, list):
            result = []
            for item in data:
                if isinstance(item, dict):
                    # 🔥 提取附件内容 🔥
                    attachment_content = self._extract_attachments_content(item, file_path)
                    if attachment_content:
                        # 将附件内容合并到现有内容中
                        if 'content' in item and isinstance(item['content'], dict):
                            if '正文内容' in item['content']:
                                item['content']['正文内容'] += f"\n\n--- 附件内容 ---\n{attachment_content}"
                            else:
                                item['content']['附件内容'] = attachment_content
                        else:
                            item['附件内容'] = attachment_content
                    
                    result.append(item)
                else:
                    result.append({"__file_source": file_path, "value": item})
            return result
            
        elif isinstance(data, dict):
            # 🔥 提取附件内容 🔥
            attachment_content = self._extract_attachments_content(data, file_path)
            if attachment_content:
                # 将附件内容合并到现有内容中
                if 'content' in data and isinstance(data['content'], dict):
                    if '正文内容' in data['content']:
                        data['content']['正文内容'] += f"\n\n--- 附件内容 ---\n{attachment_content}"
                    else:
                        data['content']['附件内容'] = attachment_content
                else:
                    data['附件内容'] = attachment_content
            
            return [data]
        else:
            return [{"__file_source": file_path, "value": data}]
    
    def consolidate_data_for_metric(self, metric_name, data_sources=None):
        """为特定指标整合数据，支持优先级"""
        # 如果没有指定数据源，使用配置中的优先级列表
        if data_sources is None:
            data_sources = METRIC_DATA_MAPPING.get(metric_name, [])
            
        if not data_sources:
            print(f"  错误: 未找到指标 '{metric_name}' 对应的数据源配置")
            return None
        
        print(f"  正在为指标 '{metric_name}' 整合数据源: {data_sources}")
        
        # 🔥 使用动态目录路径 🔥
        consolidated_dir = get_consolidated_dir_path(self.data_dir)
        dir_mode = "智能截取模式" if is_school_extraction_enabled() else "传统模式"
        print(f"  📂 使用整合目录: {consolidated_dir} ({dir_mode})")

        os.makedirs(consolidated_dir, exist_ok=True)
        
        # 🔥 文件名包含模式标识 🔥
        mode_suffix = "_intelligent" if is_school_extraction_enabled() else "_traditional"
        metric_data_file = os.path.join(consolidated_dir, f"{metric_name}{mode_suffix}_data.json")

        # 🔥 检查是否已存在相同模式的文件 🔥
        if os.path.exists(metric_data_file):
            file_age = os.path.getmtime(metric_data_file)
            from datetime import datetime
            age_str = datetime.fromtimestamp(file_age).strftime("%Y-%m-%d %H:%M:%S")
            print(f"  ℹ️  发现已存在的文件: {os.path.basename(metric_data_file)} (修改时间: {age_str})")
            
            # 可以选择是否重新生成
            # 这里默认重新生成，您也可以添加参数控制
        
        # 按优先级收集数据
        all_metric_data = []
        successful_sources = []
        
        for data_source in data_sources:
            source_data = self.read_data_from_source(data_source)
            if source_data:
                all_metric_data.extend(source_data)
                successful_sources.append(data_source)
                print(f"    ✅ 成功从 {data_source} 获取 {len(source_data)} 条数据")
            else:
                print(f"    ⚠️  数据源 {data_source} 暂无数据")
        
        if not all_metric_data:
            print(f"  警告: 未能从任何数据源 {data_sources} 中读取到数据")
            # 创建空的数据文件，避免后续处理失败
            empty_data = {
                "metric": metric_name,
                "processing_mode": "intelligent" if is_school_extraction_enabled() else "traditional",
                "data_sources": data_sources,
                "successful_sources": successful_sources,
                "total_items": 0,
                "results": [],
                "status": "no_data_found",
                "generated_at": self._get_current_time()
            }
            try:
                output_config = get_output_config(self.data_dir)
                with open(metric_data_file, 'w', encoding='utf-8') as f:
                    json.dump(empty_data, f, ensure_ascii=False, indent=output_config['json_indent'])
                return metric_data_file
            except Exception as e:
                print(f"  保存空数据文件失败: {e}")
                return None
        
        # 保存整合后的数据
        consolidated_data = {
            "metric": metric_name,
            "processing_mode": "intelligent" if is_school_extraction_enabled() else "traditional",  # 🔥 新增：标识处理模式 🔥
            "data_sources": data_sources,
            "successful_sources": successful_sources,
            "total_items": len(all_metric_data),
            "results": all_metric_data,
            "status": "success",
            "generated_at": self._get_current_time(),  # 🔥 新增：生成时间 🔥
            "attachment_extraction_enabled": is_school_extraction_enabled()  # 🔥 新增：记录附件处理模式 🔥
        }
        
        try:
            output_config = get_output_config(self.data_dir)
            with open(metric_data_file, 'w', encoding='utf-8') as f:
                json.dump(consolidated_data, f, ensure_ascii=False, indent=output_config['json_indent'])
            print(f"  ✅ 成功整合数据到: {metric_data_file}")
            print(f"  📊 数据统计: {len(all_metric_data)} 条记录，处理模式: {dir_mode}")
            return metric_data_file
        except Exception as e:
            print(f"  保存整合数据失败: {e}")
            return None
        
    def find_existing_consolidated_file(self, metric_name):
        """🔥 新增：查找已存在的整合文件（优先当前模式，后备其他模式）🔥"""
        # 当前模式的文件
        current_consolidated_dir = get_consolidated_dir_path(self.data_dir)
        current_mode_suffix = "_intelligent" if is_school_extraction_enabled() else "_traditional"
        current_file = os.path.join(current_consolidated_dir, f"{metric_name}{current_mode_suffix}_data.json")
        
        if os.path.exists(current_file):
            print(f"  ✅ 找到当前模式的文件: {os.path.basename(current_file)}")
            return current_file
        
        # 查找其他模式的文件作为备选
        other_mode = "traditional" if is_school_extraction_enabled() else "intelligent"
        other_dir = os.path.join(self.data_dir, f"consolidated_{other_mode}" if other_mode == "intelligent" else "consolidated")
        other_mode_suffix = "_traditional" if current_mode_suffix == "_intelligent" else "_intelligent"
        other_file = os.path.join(other_dir, f"{metric_name}{other_mode_suffix}_data.json")
        
        if os.path.exists(other_file):
            print(f"  ⚠️  当前模式文件不存在，找到其他模式文件: {os.path.basename(other_file)}")
            print(f"  💡 建议重新生成当前模式的数据文件")
            return other_file
        
        # 查找不带模式后缀的旧文件
        legacy_files = [
            os.path.join(current_consolidated_dir, f"{metric_name}_data.json"),
            os.path.join(self.data_dir, "consolidated", f"{metric_name}_data.json")
        ]
        
        for legacy_file in legacy_files:
            if os.path.exists(legacy_file):
                print(f"  ⚠️  找到旧格式文件: {os.path.basename(legacy_file)}")
                print(f"  💡 建议重新生成以使用新的处理模式")
                return legacy_file
        
        print(f"  ❌ 未找到指标 '{metric_name}' 的任何整合文件")
        return None
    
    def get_consolidated_file_info(self, metric_name):
        """🔥 新增：获取整合文件的详细信息 🔥"""
        current_file = self.find_existing_consolidated_file(metric_name)
        if not current_file:
            return None
        
        try:
            with open(current_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # 获取文件信息
            file_stat = os.stat(current_file)
            from datetime import datetime
            
            info = {
                'file_path': current_file,
                'file_name': os.path.basename(current_file),
                'file_size': file_stat.st_size,
                'modified_time': datetime.fromtimestamp(file_stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                'processing_mode': data.get('processing_mode', 'unknown'),
                'attachment_extraction': data.get('attachment_extraction_enabled', 'unknown'),
                'total_items': data.get('total_items', 0),
                'status': data.get('status', 'unknown'),
                'data_sources': data.get('successful_sources', [])
            }
            
            return info
            
        except Exception as e:
            print(f"  ❌ 读取文件信息失败: {e}")
            return None
    
    def list_all_consolidated_files(self):
        """🔥 新增：列出所有整合文件 🔥"""
        all_files = {}
        
        # 检查两个目录
        directories = [
            ('traditional', os.path.join(self.data_dir, 'consolidated')),
            ('intelligent', os.path.join(self.data_dir, 'consolidated_intelligent'))
        ]
        
        for mode, dir_path in directories:
            if os.path.exists(dir_path):
                for filename in os.listdir(dir_path):
                    if filename.endswith('_data.json'):
                        file_path = os.path.join(dir_path, filename)
                        
                        # 提取指标名称
                        if filename.endswith('_traditional_data.json'):
                            metric_name = filename.replace('_traditional_data.json', '')
                        elif filename.endswith('_intelligent_data.json'):
                            metric_name = filename.replace('_intelligent_data.json', '')
                        else:
                            metric_name = filename.replace('_data.json', '')
                        
                        if metric_name not in all_files:
                            all_files[metric_name] = {}
                        
                        # 获取文件信息
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                data = json.load(f)
                            
                            file_stat = os.stat(file_path)
                            from datetime import datetime
                            
                            all_files[metric_name][mode] = {
                                'file_path': file_path,
                                'file_name': filename,
                                'total_items': data.get('total_items', 0),
                                'status': data.get('status', 'unknown'),
                                'modified_time': datetime.fromtimestamp(file_stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                            }
                        except:
                            all_files[metric_name][mode] = {
                                'file_path': file_path,
                                'file_name': filename,
                                'error': 'Failed to read file'
                            }
        
        return all_files
    
    def extract_text_content(self, data_file_path):
        """从JSON文件中提取文本内容供LLM使用"""
        try:
            with open(data_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            text_list = []
            
            def extract_strings(value, path=""):
                if isinstance(value, str):
                    # 🔥 过滤文件路径，保留所有其他字符串 🔥
                    if not (len(value) > 20 and ('\\' in value or '/' in value) and value.endswith('.json')):
                        text_list.append(value)
                elif isinstance(value, (int, float)):
                    # 🔥 关键修复：包含所有数字数据 🔥
                    text_list.append(str(value))
                elif isinstance(value, list):
                    for i, item in enumerate(value):
                        extract_strings(item, f"{path}[{i}]")
                elif isinstance(value, dict):
                    for k, v in value.items():
                        if k != '__file_source':
                            # 🔥 关键修复：保留字段名称，便于LLM理解数据结构 🔥
                            if isinstance(v, (str, int, float)):
                                text_list.append(f"{k}: {v}")
                            else:
                                text_list.append(f"{k}:")  # 添加字段名
                                extract_strings(v, f"{path}.{k}" if path else k)
            
            if 'results' in data:
                print(f"    📊 提取results中的内容，共{len(data['results'])}项")
                extract_strings(data['results'])
            else:
                print(f"    📊 提取整个数据结构")
                extract_strings(data)
            
            extracted_text = "\n".join(text_list)
            print(f"    📄 提取到文本长度: {len(extracted_text)} 字符")
            
            return extracted_text
            
        except Exception as e:
            print(f"    ❌ 提取文本内容失败: {e}")
            return ""
    
    def get_data_source_info(self):
        """获取数据源信息和状态"""
        info = {
            'configured_sources': len(DATA_SOURCES),
            'existing_sources': 0,
            'missing_sources': [],
            'source_details': {}
        }
        
        for source_name, source_path in DATA_SOURCES.items():
            # 🔥 修正路径计算 🔥
            if source_path.startswith('../data/'):
                clean_path = source_path.replace('../data/', '')
            else:
                clean_path = source_path
            
            full_path = os.path.join(self.data_dir, clean_path)
            exists = os.path.exists(full_path)
            
            if exists:
                info['existing_sources'] += 1
                # 统计该数据源中的文件数量
                file_count = 0
                if os.path.isdir(full_path):
                    for root, _, files in os.walk(full_path):
                        file_count += len([f for f in files if f.endswith('.json')])
                
                info['source_details'][source_name] = {
                    'path': full_path,
                    'exists': True,
                    'file_count': file_count
                }
            else:
                info['missing_sources'].append(source_name)
                info['source_details'][source_name] = {
                    'path': full_path,
                    'exists': False,
                    'file_count': 0
                }
        
        return info
    
    def _extract_attachments_content(self, data_item, file_path):
        """提取数据项中的附件内容 - 支持可选的智能截取模式"""
        # 检查是否启用附件处理
        attachment_config = get_attachment_config()
        if not attachment_config.get('enable_attachment_processing', True):
            print(f"    ⚠️  附件处理已禁用")
            return ""
        
        attachment_contents = []
        
        if 'content' in data_item and '附件' in data_item['content']:
            attachments = data_item['content']['附件']
            if isinstance(attachments, list):
                base_path = os.path.dirname(file_path)
                
                for attachment in attachments:
                    if isinstance(attachment, dict) and attachment.get('download_status') == 'success':
                        local_path = attachment.get('local_path', '')
                        if local_path:
                            full_path = os.path.join(base_path, local_path)
                            
                            if os.path.exists(full_path):
                                file_ext = os.path.splitext(full_path)[1].lower()
                                attachment_name = attachment.get('name', '未知附件')
                                
                                # 检查文件格式是否支持
                                if file_ext not in attachment_config.get('supported_formats', []):
                                    print(f"    ⚠️  不支持的附件格式: {file_ext}")
                                    continue
                                
                                # 检查文件大小
                                try:
                                    file_size = os.path.getsize(full_path)
                                    max_size = attachment_config.get('max_attachment_size', 50 * 1024 * 1024)
                                    if file_size > max_size:
                                        print(f"    ⚠️  附件过大: {file_size} bytes (限制: {max_size} bytes)")
                                        continue
                                except:
                                    pass
                                
                                try:
                                    # 读取原始内容
                                    raw_content = self._read_attachment_by_type(full_path, file_ext)
                                    
                                    if raw_content:
                                        # 🔥 根据配置选择处理模式 🔥
                                        if is_school_extraction_enabled():
                                            print(f"    📍 使用智能截取模式处理附件: {attachment_name}")
                                            extracted_content = self._extract_school_relevant_content(raw_content, attachment_name)
                                        else:
                                            print(f"    📄 使用传统模式处理附件: {attachment_name}")
                                            extracted_content = self._extract_traditional_content(raw_content, attachment_name)
                                        
                                        if extracted_content and len(extracted_content) > 20:
                                            formatted_content = f"\n=== 附件：{attachment_name} ===\n{extracted_content}\n=== 附件结束 ===\n"
                                            attachment_contents.append(formatted_content)
                                            print(f"    ✅ 成功提取附件内容: {attachment_name} ({len(extracted_content)} 字符)")
                                        else:
                                            print(f"    ⚠️  附件内容为空或过短: {attachment_name}")
                                    else:
                                        print(f"    ⚠️  无法读取附件内容: {attachment_name}")
                                        
                                except Exception as e:
                                    print(f"    ❌ 处理附件 {attachment_name} 时出错: {e}")
                            else:
                                print(f"    ⚠️  附件文件不存在: {full_path}")
        
        return "\n".join(attachment_contents)
    
    def _read_attachment_by_type(self, file_path, file_ext):
        """根据文件类型读取附件内容"""
        try:
            if file_ext == '.pdf':
                return self._read_pdf_attachment(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._read_docx_attachment(file_path)
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                print(f"    ⚠️  不支持的文件类型: {file_ext}")
                return ""
        except Exception as e:
            print(f"    ❌ 读取文件 {file_path} 失败: {e}")
            return ""
    
    def _extract_traditional_content(self, raw_content, attachment_name):
        """传统模式：提取附件的完整内容（带长度限制和清理）"""
        if not raw_content or not isinstance(raw_content, str):
            return ""
        
        print(f"    📄 传统模式处理附件: {attachment_name}")
        
        # 获取传统处理配置
        traditional_config = get_traditional_extraction_config()
        
        # 基础清理
        cleaned_content = raw_content
        
        # 清理HTML标签（如果启用）
        if traditional_config.get('clean_html_tags', True):
            cleaned_content = self._remove_html_tags(cleaned_content)
        
        # 清理多余空白（如果启用）
        if traditional_config.get('remove_extra_whitespace', True):
            cleaned_content = self._basic_clean_text(cleaned_content)
        
        # 限制内容长度
        max_length = traditional_config.get('max_content_length', 10000)
        if len(cleaned_content) > max_length:
            cleaned_content = cleaned_content[:max_length]
            print(f"    ✂️  内容过长，截取前 {max_length} 字符")
        
        return cleaned_content
    
    def _remove_html_tags(self, text):
        """移除HTML标签"""
        if not text:
            return ""
        
        import re
        # 移除HTML标签
        clean_text = re.sub(r'<[^>]+>', '', text)
        # 解码HTML实体
        try:
            from html import unescape
            clean_text = unescape(clean_text)
        except:
            pass
        
        return clean_text
    
    def _extract_school_relevant_content(self, raw_content, attachment_name):
        """基于学校名称提取相关内容片段 - 支持配置化"""
        if not raw_content or not isinstance(raw_content, str):
            return ""
        
        # 获取学校截取配置
        school_config = get_school_extraction_config()
        
        # 获取学校列表
        school_list = self._get_school_list(school_config.get('school_list_source', 'csv'))
        
        if not school_list:
            print(f"    ⚠️  未能获取学校列表，切换到传统模式")
            return self._extract_traditional_content(raw_content, attachment_name)
        
        print(f"    🔍 智能截取模式：在附件 {attachment_name} 中搜索 {len(school_list)} 个学校的相关内容...")
        
        # 清理原始内容
        cleaned_content = self._basic_clean_text(raw_content)
        
        extracted_segments = []
        found_schools = []
        
        for school_name in school_list:
            # 查找学校相关片段
            school_segments = self._find_school_segments_configurable(cleaned_content, school_name, school_config)
            if school_segments:
                found_schools.append(school_name)
                extracted_segments.extend(school_segments)
                print(f"    📍 找到 {school_name}，共 {len(school_segments)} 处")
        
        if extracted_segments:
            print(f"    ✅ 找到 {len(found_schools)} 个学校的相关内容：{', '.join(found_schools[:5])}{'...' if len(found_schools) > 5 else ''}")
            
            # 按位置排序并构建最终内容
            extracted_segments.sort(key=lambda x: x['position'])
            
            final_content_parts = []
            for segment in extracted_segments:
                final_content_parts.append(
                    f"📍 {segment['school']}：{segment['content']}..."
                )
            
            return "\n\n".join(final_content_parts)
        else:
            print(f"    ⚠️  未找到任何目标学校，切换到传统模式")
            return self._extract_traditional_content(raw_content, attachment_name)
    
    def _find_school_segments_configurable(self, content, school_name, config):
        """根据配置查找学校片段"""
        segments = []
        
        if not content or not school_name:
            return segments
        
        # 获取配置参数
        chars_after = config.get('characters_after_school', 100)
        chars_before = config.get('characters_before_school', 0)
        max_segments = config.get('max_segments_per_school', 10)
        min_length = config.get('min_segment_length', 10)
        
        start_pos = 0
        segment_count = 0
        
        while segment_count < max_segments:
            pos = content.find(school_name, start_pos)
            if pos == -1:
                break
            
            # 计算片段范围
            school_start_pos = pos
            school_end_pos = pos + len(school_name)
            
            segment_start = max(0, school_start_pos - chars_before)
            segment_end = min(len(content), school_end_pos + chars_after)
            
            # 提取片段
            segment_content = content[segment_start:segment_end]
            cleaned_segment = self._clean_segment_text(segment_content)
            
            if cleaned_segment and len(cleaned_segment.strip()) >= min_length:
                segments.append({
                    'school': school_name,
                    'position': pos,
                    'content': cleaned_segment,
                    'start': segment_start,
                    'end': segment_end
                })
                segment_count += 1
            
            start_pos = pos + len(school_name)
        
        return segments
    
    def _get_school_list(self, source_type='csv'):
        """获取学校列表 - 修复路径计算"""
        school_list = []
        
        print(f"    🔍 尝试获取学校列表，来源类型: {source_type}")
        
        try:
            if source_type == 'csv':
                # 🔥 修复CSV文件路径计算 🔥
                # 当前文件: scrapetest/src/AIquest/utils/data_reader.py
                # CSV文件: scrapetest/ai_evaluation_dataset_long.csv
                
                current_file_dir = os.path.dirname(os.path.abspath(__file__))  # utils目录
                aiquest_dir = os.path.dirname(current_file_dir)  # AIquest目录  
                src_dir = os.path.dirname(aiquest_dir)  # src目录
                project_root = os.path.dirname(src_dir)  # scrapetest目录 (项目根目录)
                
                csv_file = os.path.join(project_root, "ai_evaluation_dataset_long.csv")
                
                print(f"    📂 路径计算详情:")
                print(f"      当前文件: {os.path.abspath(__file__)}")
                print(f"      utils目录: {current_file_dir}")
                print(f"      AIquest目录: {aiquest_dir}")
                print(f"      src目录: {src_dir}")
                print(f"      项目根目录: {project_root}")
                print(f"      CSV文件路径: {csv_file}")
                print(f"      CSV文件存在: {os.path.exists(csv_file)}")
                
                # 🔥 如果主路径不存在，尝试备选路径 🔥
                if not os.path.exists(csv_file):
                    print(f"    🔍 主路径不存在，尝试备选路径:")
                    alternative_paths = [
                        # 可能在项目根目录的不同位置
                        os.path.join(project_root, "data", "ai_evaluation_dataset_long.csv"),
                        os.path.join(src_dir, "ai_evaluation_dataset_long.csv"),
                        os.path.join(src_dir, "data", "ai_evaluation_dataset_long.csv"),
                        # 可能在上级目录
                        os.path.join(os.path.dirname(project_root), "ai_evaluation_dataset_long.csv"),
                        # 绝对路径（如果您确定文件在特定位置）
                        r"c:\Users\83789\PycharmProjects\scrapetest\ai_evaluation_dataset_long.csv"
                    ]
                    
                    for i, alt_path in enumerate(alternative_paths, 1):
                        exists = os.path.exists(alt_path)
                        print(f"      备选路径{i}: {alt_path} ({'✅存在' if exists else '❌不存在'})")
                        if exists:
                            csv_file = alt_path
                            break
                
                # 如果找到CSV文件，尝试读取
                if os.path.exists(csv_file):
                    try:
                        import csv
                        with open(csv_file, 'r', encoding='utf-8') as f:
                            reader = csv.DictReader(f)
                            schools = set()
                            row_count = 0
                            
                            # 读取前几行来调试
                            sample_rows = []
                            for row in reader:
                                row_count += 1
                                if row_count <= 3:
                                    sample_rows.append(dict(row))
                                
                                school_name = row.get('学校名称', '').strip()
                                if school_name and school_name != '学校名称' and school_name != '':
                                    schools.add(school_name)
                            
                            school_list = list(schools)
                            
                            print(f"    📊 CSV文件读取结果:")
                            print(f"      总行数: {row_count}")
                            print(f"      唯一学校数: {len(school_list)}")
                            print(f"      前3行示例: {sample_rows}")
                            print(f"      前5个学校: {school_list[:5] if school_list else '无'}")
                            
                            if school_list:
                                return school_list
                            else:
                                print(f"    ⚠️  CSV文件中没有找到有效的学校名称")
                                
                    except Exception as csv_error:
                        print(f"    ❌ 读取CSV文件出错: {csv_error}")
                        import traceback
                        print(f"    错误详情: {traceback.format_exc()}")
                else:
                    print(f"    ❌ 所有路径都找不到CSV文件")
            
            # 如果CSV失败，使用预定义列表
            if source_type == 'predefined' or not school_list:
                print(f"    📋 使用预定义学校列表")
                school_list = [
                    '中山大学', '暨南大学', '华南理工大学', '华南农业大学', '广州医科大学',
                    '广州中医药大学', '广东药科大学', '华南师范大学', '广州体育学院', '广州美术学院',
                    '星海音乐学院', '广东技术师范大学', '广东财经大学', '广州大学', '广州航海学院',
                    '广东警官学院', '仲恺农业工程学院', '广东金融学院', '广东工业大学', '广东外语外贸大学',
                    '南方医科大学', '广东第二师范学院', '广东轻工职业技术大学', '广东白云学院', '广东培正学院',
                    '广州城市理工学院', '广州软件学院', '广州南方学院', '广东外语外贸大学南国商学院',
                    '广州华商学院', '华南农业大学珠江学院', '广州理工学院', '广州华立学院', '广州应用科技学院',
                    '广州商学院', '广州工商学院', '广州科技职业技术大学', '广州新华学院', '香港科技大学（广州）',
                    '广州番禺职业技术学院'
                ]
                print(f"    ✅ 预定义学校列表加载成功，共 {len(school_list)} 个学校")
                return school_list
            
            elif source_type == 'auto':
                # 自动模式：先尝试CSV，失败则用预定义
                csv_result = self._get_school_list('csv')
                if csv_result:
                    return csv_result
                return self._get_school_list('predefined')
        
        except Exception as e:
            print(f"    ❌ 获取学校列表出错: {e}")
            import traceback
            print(f"    错误详情: {traceback.format_exc()}")
        
        # 最后的备选方案
        if not school_list:
            school_list = ['中山大学', '暨南大学', '华南理工大学', '华南农业大学', '华南师范大学']
            print(f"    🚑 使用紧急备选学校列表，共 {len(school_list)} 个学校")
        
        return school_list
    
    def _find_school_segments(self, content, school_name):
        """🔥 已弃用：使用 _find_school_segments_configurable 替代 🔥"""
        # 使用默认配置调用新方法
        default_config = {
            'characters_after_school': 100,
            'characters_before_school': 0,
            'max_segments_per_school': 10,
            'min_segment_length': 10
        }
        return self._find_school_segments_configurable(content, school_name, default_config)
    
    def _basic_clean_text(self, text):
        """基础文本清理，保留基本格式"""
        if not text or not isinstance(text, str):
            return ""
        
        import re
        
        # 移除过多的换行符和空白
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\t+', ' ', text)
        
        # 移除特殊控制字符，但保留基本标点
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
        
        # 清理多余空格，但保留单个空格
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def _clean_segment_text(self, segment_text):
        """清理片段文本"""
        if not segment_text or not isinstance(segment_text, str):
            return ""
        
        import re
        
        # 基础清理
        segment_text = self._basic_clean_text(segment_text)
        
        # 移除可能的文件路径和URL
        segment_text = re.sub(r'[A-Za-z]:\\[^\s]+', '', segment_text)
        segment_text = re.sub(r'https?://[^\s]+', '', segment_text)
        
        # 移除过长的数字串（可能是ID或无意义数据）
        segment_text = re.sub(r'\b\d{10,}\b', '', segment_text)
        
        # 清理多余的标点符号
        segment_text = re.sub(r'[。，]{3,}', '...', segment_text)
        segment_text = re.sub(r'[-_=]{5,}', '---', segment_text)
        
        # 最终清理
        segment_text = re.sub(r'\s+', ' ', segment_text)
        
        return segment_text.strip()
    
    def _clean_attachment_text(self, text):
        """清理附件文本中的格式字符 - 保持原有功能"""
        if not text or not isinstance(text, str):
            return ""
        
        import re
        
        # 移除换行符和特殊字符
        text = re.sub(r'[\n\r\t]+', '', text)
        text = re.sub(r'\s+', ' ', text)
        
        # 移除转义字符
        text = re.sub(r'\\[nrt]', ' ', text)
        
        # 清理Unicode特殊字符
        text = re.sub(r'[\u00a0\u2000-\u200f\u2028-\u202f\u3000]', ' ', text)
        
        # 保留中文和常用标点
        text = re.sub(r'[^\w\s\u4e00-\u9fff\u3001.,;:：!?()[\]{}"\'-]', '', text)
        text = text.replace('、', ', ')
        
        return text.strip()
    
    def _read_pdf_attachment(self, pdf_path):
        """读取PDF附件内容"""
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_content = ""
                for page in reader.pages:
                    text_content += page.extract_text() + " "
                return text_content
        except ImportError:
            print(f"    ⚠️  PyPDF2未安装，无法读取PDF")
            return ""
        except Exception as e:
            print(f"    ❌ 读取PDF时出错: {e}")
            return ""
    
    def _read_docx_attachment(self, docx_path):
        """读取DOCX附件内容"""
        try:
            import docx
            doc = docx.Document(docx_path)
            text_content = ""
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + " "
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
            
            return text_content
        except ImportError:
            print(f"    ⚠️  python-docx未安装，无法读取DOCX")
            return ""
        except Exception as e:
            print(f"    ❌ 读取DOCX时出错: {e}")
            return ""